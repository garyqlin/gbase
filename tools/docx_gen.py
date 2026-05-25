# SPDX-License-Identifier: MIT
"""Word document generator"""
import json
import logging
import os
import subprocess

from lib.toolkit import tool

logger = logging.getLogger(__name__)

@tool()
async def gen_docx(title: str = "文档", content: list = None, output_path: str = "") -> dict:
    """
    生成 Word (.docx) 文档。
    
    Args:
        title: 文档标题
        content: 内容结构，每项为 {"type": "h1|h2|h3|p|table|image", "text": "...", ...}
                  - h1/h2/h3: {"type": "h1", "text": "标题"}
                  - p: {"type": "p", "text": "段落文字"}
                  - table: {"type": "table", "headers": ["列1","列2"], "rows": [["a","b"]]}
                  - image: {"type": "image", "path": "/path/to/img.png"}
        output_path: 输出路径，为空则自动生成在 ~/Downloads/
    
    Returns:
        {"path": "...", "size": 12345}
    """
    if content is None:
        content = [{"type": "p", "text": "（空文档）"}]

    if not output_path:
        output_path = os.path.expanduser(f"~/Downloads/{title}.docx")

    script = f'''import docx
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

doc = docx.Document()
# 标题
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'

doc.add_heading({json.dumps(title)}, level=0)

content = {json.dumps(content)}
for item in content:
    t = item.get("type", "p")
    if t == "h1":
        doc.add_heading(item["text"], level=1)
    elif t == "h2":
        doc.add_heading(item["text"], level=2)
    elif t == "h3":
        doc.add_heading(item["text"], level=3)
    elif t == "p":
        p = doc.add_paragraph(item["text"])
        p.paragraph_format.space_after = Pt(6)
    elif t == "table":
        headers = item.get("headers", [])
        rows = item.get("rows", [])
        table = doc.add_table(rows=1+len(rows), cols=max(len(headers), 1))
        table.style = 'Light Shading Accent 1'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci < len(table.rows[ri+1].cells):
                    table.rows[ri+1].cells[ci].text = str(val)
        doc.add_paragraph()
    elif t == "image":
        path = item.get("path", "")
        if path and os.path.exists(path):
            doc.add_picture(path, width=Inches(5.5))

doc.save({json.dumps(output_path)})
print("OK:" + {json.dumps(output_path)})
'''

    result = subprocess.run(
        ["python3", "-c", script],
        capture_output=True, text=True, timeout=30
    )

    if result.returncode != 0:
        return {"error": result.stderr.strip() or result.stdout.strip()}

    if os.path.isfile(output_path):
        size = os.path.getsize(output_path)
        return {"path": output_path, "size": size}
    return {"error": "文件未生成"}
